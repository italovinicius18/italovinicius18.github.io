import os
import yaml
from jinja2 import Environment, FileSystemLoader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document


class YamlParser:
    def __init__(self, yaml_file):
        self.yaml_file = yaml_file

    def parse_yaml(self):
        with open(self.yaml_file, "r", encoding="utf-8") as file:
            data = yaml.safe_load(file)
        return data

    def generate_html(self, template_file, output_file):
        data = self.parse_yaml()
        env = Environment(loader=FileSystemLoader(os.path.join("src", "jinja")))
        template = env.get_template(template_file)

        # Read styles from CSS file
        with open(os.path.join("src", "jinja", "styles.css"), "r") as file:
            styles = file.read()

        html_content = template.render(data=data, styles=styles)

        output_path = os.path.join(output_file)
        with open(output_path, "w", encoding="utf-8") as file:
            file.write(html_content)

    def generate_markdown(self, output_file):
        data = self.parse_yaml()

        with open(os.path.join("dist", output_file), "w", encoding="utf-8") as file:
            for key, value in data.items():
                file.write(f"## {key}\n")
                if isinstance(value, list):
                    for item in value:
                        file.write(f"- {item}\n")
                else:
                    file.write(f"- {value}\n")

    def generate_pdf(self, html_file, output_file):
        from xhtml2pdf import pisa

        with open(
            os.path.join("dist", html_file), "r", encoding="utf-8"
        ) as source_file:
            html_content = source_file.read()

        pdf_path = os.path.join("dist", output_file)
        pdf_file = open(pdf_path, "wb")
        pisa_status = pisa.CreatePDF(html_content, dest=pdf_file)
        pdf_file.close()

    def generate_docx(self, output_file):
        data = self.parse_yaml()
        doc = Document()

        for key, value in data.items():
            doc.add_heading(key, level=1)
            if isinstance(value, list):
                for item in value:
                    doc.add_paragraph(item)
            else:
                doc.add_paragraph(value)

        doc_path = os.path.join("dist", output_file)
        doc.save(doc_path)


if __name__ == "__main__":
    yaml_parser = YamlParser(os.path.join("resume.yaml"))

    # Generate HTML
    yaml_parser.generate_html("template.html", "index.html")

    # # Generate Markdown
    # yaml_parser.generate_markdown("resume.md")

    # # Generate PDF from HTML
    # yaml_parser.generate_pdf("resume.html", "resume.pdf")

    # Generate DOCX
    # yaml_parser.generate_docx("resume.docx")  
